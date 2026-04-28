"""
make_word.py
------------
Renders REPORT.md to REPORT.docx for Gradescope submission.
Mirrors the structure of make_report.py (the PDF generator) so the Word
file reads the same: title block, headings, body paragraphs, embedded
charts, monospace inline code.

Run:  python3 make_word.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).parent
SRC = HERE / "REPORT.md"
OUT = HERE / "REPORT.docx"


def _set_run_color(run, hex_color):
    run.font.color.rgb = RGBColor.from_string(hex_color.lstrip("#"))


def _add_inline(paragraph, text):
    """Parse minimal markdown inline syntax and add styled runs."""
    # Tokenize: split on **bold**, *italic*, `code`, keeping delimiters
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6,
                     space_before=0):
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)


def render(md: str, doc: Document):
    lines = md.splitlines()
    i = 0
    paragraph_buffer = []

    def flush_paragraph():
        if paragraph_buffer:
            text = " ".join(paragraph_buffer).strip()
            paragraph_buffer.clear()
            if not text:
                return
            p = doc.add_paragraph()
            _style_paragraph(p, space_after=6)
            _add_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(11)

    while i < len(lines):
        line = lines[i].rstrip()

        # Image
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if m:
            flush_paragraph()
            caption, path = m.group(1), m.group(2)
            img_path = HERE / path
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(5.8))
                if caption:
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.space_after = Pt(10)
                    cap_run = cap_p.add_run(caption)
                    cap_run.italic = True
                    cap_run.font.size = Pt(9)
                    _set_run_color(cap_run, "#555555")
            i += 1
            continue

        # Title
        if line.startswith("# "):
            flush_paragraph()
            p = doc.add_paragraph()
            _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(20)
            run.bold = True
            _set_run_color(run, "#1a1a2e")
            i += 1
            continue

        # H2
        if line.startswith("## "):
            flush_paragraph()
            p = doc.add_paragraph()
            _style_paragraph(p, space_before=14, space_after=6)
            run = p.add_run(line[3:].strip())
            run.font.size = Pt(14)
            run.bold = True
            _set_run_color(run, "#1a1a2e")
            i += 1
            continue

        # H3
        if line.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph()
            _style_paragraph(p, space_before=10, space_after=4)
            run = p.add_run(line[4:].strip())
            run.font.size = Pt(11.5)
            run.bold = True
            _set_run_color(run, "#2c3e7a")
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            flush_paragraph()
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "cccccc")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Author / date subtitle (bolded line near top)
        if line.startswith("**Authors:**") or line.startswith("**Date:**"):
            flush_paragraph()
            p = doc.add_paragraph()
            _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
            _add_inline(p, line)
            for run in p.runs:
                run.font.size = Pt(11)
                _set_run_color(run, "#444444")
            i += 1
            continue

        # Bold-only subtitle line (like "**CIS 1921 Final Project Report**")
        if line.startswith("**") and line.endswith("**") and "**" not in line[2:-2]:
            flush_paragraph()
            p = doc.add_paragraph()
            _style_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
            run = p.add_run(line.strip("*").strip())
            run.font.size = Pt(11)
            _set_run_color(run, "#444444")
            i += 1
            continue

        # Blank line ends a paragraph
        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        # Otherwise accumulate
        paragraph_buffer.append(line.strip())
        i += 1

    flush_paragraph()


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()

    # Margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    render(md, doc)
    doc.save(str(OUT))
    print(f"Wrote {OUT} ({OUT.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
